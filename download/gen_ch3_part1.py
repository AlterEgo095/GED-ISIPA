#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate Chapter 3 DOCX for GED-ISIPA Master's Thesis
Part 1: Imports, helpers, and setup
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ─── PATHS ───
DIAGRAMS_DIR = "/home/z/my-project/download/ch3_diagrams"
SCREENSHOTS_DIR = "/home/z/my-project/download/screenshots"
OUTPUT_PATH = "/home/z/my-project/download/Chapitre3_GED_ISIPA_Definitif.docx"

# ─── COLORS ───
DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
MEDIUM_BLUE = RGBColor(0x2C, 0x5F, 0x8A)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "1B3A5C"
TABLE_ALT_BG = "EBF0F7"

# ─── FIGURE COUNTER ───
fig_counter = [0]

def next_fig():
    fig_counter[0] += 1
    return fig_counter[0]

# ─── HELPER FUNCTIONS ───

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = DARK_BLUE
            run.font.bold = True
            run.font.name = 'Cambria'
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = MEDIUM_BLUE
            run.font.bold = True
            run.font.name = 'Cambria'
        elif level == 3:
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_GRAY
            run.font.bold = True
            run.font.name = 'Cambria'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria')
    return h

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(3)
    pf.line_spacing = 1.5
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Cambria'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria')
    return p

def add_body_rich(doc, parts):
    """parts is a list of (text, bold, italic) tuples"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(3)
    pf.line_spacing = 1.5
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Cambria'
        run.bold = bold
        run.italic = italic
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria')
    return p

def add_image_with_caption(doc, img_path, caption_text, width_cm=15):
    fig_num = next_fig()
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
    else:
        add_body(doc, f"[Image non trouvée: {img_path}]")

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(f"Figure 3.{fig_num} : {caption_text}")
    cap_run.font.size = Pt(10)
    cap_run.font.italic = True
    cap_run.font.bold = True
    cap_run.font.name = 'Cambria'
    cap_run.font.color.rgb = DARK_GRAY
    cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria')
    return fig_num

def add_analysis(doc, paragraphs):
    for text in paragraphs:
        add_body(doc, text)

def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.font.name = 'Cambria'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria')
        set_cell_shading(cell, TABLE_HEADER_BG)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = 'Cambria'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria')
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Cambria'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria')
    return p

def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Cambria'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria')
    return p

print("Part 1 loaded: helpers ready")
