#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Chapitre 3 GED-ISIPA - Part 1: Setup, Helper Functions, and Content Sections 3.1-3.4"""
import os, sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT = '/home/z/my-project/download/Chapitre3_GED_ISIPA_Definitif.docx'
DIAG_DIR = '/home/z/my-project/download/ch3_diagrams'
SS_DIR = '/home/z/my-project/download/screenshots'

# ============ Helper Functions ============
def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F) if level <= 2 else RGBColor(0x2C, 0x3E, 0x50)
    return h

def add_para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=11, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    return p

def add_figure(doc, img_path, caption, width_cm=15):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.name = 'Calibri'
    r.bold = True
    r.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '1E3A5F')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9)
                r.font.name = 'Calibri'
                r.bold = True
    # Data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = 'Calibri'
            if ri % 2 == 1:
                set_cell_shading(cell, 'F0F4F8')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacing
    return table

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    for r in p.runs:
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
    return p

# ============ Document Setup ============
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============ CHAPTER TITLE ============
title = doc.add_heading('CHAPITRE III : IMPLEMENTATION DE LA SOLUTION GED-ISIPA', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    run.font.size = Pt(20)

doc.add_paragraph()

# ============ 3.1 PRESENTATION GENERALE ============
add_heading(doc, '3.1 Presentation generale de la solution realisee', 1)

add_para(doc, "Le present chapitre est consacre a la presentation detaillee de la solution informatique developpee dans le cadre de ce travail de fin d'etudes. Cette solution, denommee GED-ISIPA (Gestion Electronique de Documents de l'ISIPA), repond a la problematique formulee dans l'introduction de ce memoire : comment concevoir et implementer une solution numerique permettant de reduire la lenteur des operations documentaires, le risque de perte des documents, la dispersion des informations, l'absence de tracabilite et le manque de securite dans la gestion documentaire de l'ISIPA. La solution developpee s'inscrit dans la continuite des objectifs definis au Chapitre I et du cadre methodologique Merise et UML adopte pour la conception du systeme.")

add_para(doc, "La solution GED-ISIPA se distingue par sa capacite a s'adapter a differents contextes organisationnels. En effet, la plateforme prend en charge huit types d'organisations : universite, hopital, entreprise, gouvernement, PME, institution, ONG et cabinet juridique. Cette extension du perimetre initial, initialement limite a l'ISIPA, repond a une logique de mutualisation et de scalabilite : l'architecture multi-tenant permet a chaque organisation de beneficier d'une instance dediee au sein d'une plateforme commune, garantissant l'isolation totale de ses donnees tout en profitant des evolutions fonctionnelles centralisees. Cette approche SaaS (Software as a Service) constitue une valeur ajoutee significative par rapport a l'objectif initial, en transformant un outil mono-organisation en une plateforme evolutive.")

add_para(doc, "Les objectifs atteints par cette implementation peuvent etre ressumes comme suit, en correspondance avec les objectifs specifiques definis au Chapitre I. Premierement, l'analyse des insuffisances du systeme actuel a permis d'identifier cinq problematiques majeures (lenteur, risque de perte, dispersion, absence de tracabilite, manque de securite) qui ont guide la conception de la solution. Deuxiemement, les besoins fonctionnels et techniques ont ete formellement definis et traduits en une architecture logicielle adaptee, incluant un moteur de workflow, une matrice de permissions RBAC granulaire et un systeme multi-tenant. Troisiemement, l'architecture du systeme a ete concue selon les principes de la methode Merise (MCD, MLD, MPD) et modelisee a l'aide de diagrammes UML. Quatriemement, la solution a ete developpee et implementee avec les technologies modernes du stack Next.js. Cinquiemement, des tests de validation fonctionnelle ont ete realises sur l'application deployee. Sixiemement, des recommandations pour la maintenance et l'evolution du systeme sont formulees en fin de chapitre.")

add_para(doc, "La valeur ajoutee de cette solution reside dans sa capacite a unifier, au sein d'une plateforme coherente, des fonctionnalites qui etaient jusqu'alors disjointes ou inexistantes. L'ISIPA, comme de nombreuses institutions congolaises, fonctionnait avec un systeme de gestion documentaire essentiellement manuel, caracterise par des processus lents, une absence de tracabilite et des risques importants de perte ou d'acces non autorise aux documents sensibles. La solution GED-ISIPA apporte une reponse concrete a ces problemes en offrant une gestion electronique complete du cycle de vie documentaire, depuis la creation d'un brouillon jusqu'a l'archivage definitif, en passant par les etapes de revision, d'approbation et de publication.")

# ============ 3.2 ARCHITECTURE GENERALE ============
add_heading(doc, '3.2 Architecture generale du systeme', 1)

add_heading(doc, '3.2.1 Architecture logique', 2)

add_para(doc, "L'architecture logique du systeme GED-ISIPA est organisee en cinq couches distinctes, chacune assurant une responsabilite specifique et communiquant avec les couches adjacentes selon des interfaces bien definies. Cette decomposition en couches respecte les principes d'architecture logicielle presents par Sommerville (2023) et Fowler (2002), garantissant la separation des preoccupations (separation of concerns), la cohabitation evolutive et la maintenabilite du systeme. Les cinq couches sont : la couche de presentation, la couche API et metier, la couche d'authentification et de securite, la couche d'acces aux donnees et la couche infrastructure.")

add_figure(doc, f'{DIAG_DIR}/fig_archi_logique.png', 'Figure 3.1 : Architecture logique en couches du systeme GED-ISIPA')

add_para(doc, "La couche de presentation constitue l'interface utilisateur de l'application. Elle est implementee a l'aide du framework React 19 avec le systeme de routage App Router de Next.js 16. Les composants d'interface sont construits sur la bibliotheque shadcn/ui, qui fournit des composants accessibles et personnalisables bases sur les primitives Radix UI. Le style visuel est assure par Tailwind CSS 4, offrant une approche utility-first pour la conception responsive. Cette couche comprend 25 pages (page.tsx) et plus de 40 composants React organises par domaine fonctionnel : authentification, documents, workflows, modules, tableaux de bord et administration.")

add_para(doc, "La couche API et metier regroupe l'ensemble de la logique fonctionnelle de l'application. Elle est implementee sous forme de Route Handlers Next.js, exposant 41 endpoints RESTful pour les operations de creation, lecture, mise a jour et suppression (CRUD), ainsi que des operations metier specifiques comme la soumission, l'approbation, la publication et l'archivage de documents. Cette couche intègre egalement les moteurs metier : le moteur de workflow (workflow.ts), le moteur de modules (module-engine.ts), la matrice de permissions (permissions.ts), le generateur de jetons (token-engine.ts) et le moteur de redirection (redirection.ts). Chaque endpoint est protege par une validation Zod et un controle RBAC.")

add_para(doc, "La couche d'authentification et de securite assure la gestion des identites et le controle d'acces. Elle repose sur NextAuth.js en strategie JWT, avec un middleware implementant le rate limiting, l'injection d'en-tetes de securite HTTP (X-Frame-Options, X-Content-Type-Options, Referrer-Policy, X-XSS-Protection) et la verification systematique des jetons d'authentification. La protection des routes est assuree a deux niveaux : au niveau du middleware (protection coarse-grained par route) et au niveau des handlers API (protection fine-grained par permission RBAC).")

add_para(doc, "La couche d'acces aux donnees est assuree par Prisma ORM, qui fournit une abstraction type-safe au-dessus de la base de donnees. En environnement de developpement, SQLite est utilise pour sa simplicite de mise en oeuvre, tandis que PostgreSQL est configure dans l'architecture de production via Docker Compose. Le schema Prisma definit 14 modeles, 10 enumerations et les relations entre entites, garantissant la coherence des donnees et l'isolation multi-tenant via le champ organizationId present dans chaque requete.")

add_para(doc, "Enfin, la couche infrastructure englobe l'ensemble des composants de deploiement : conteneurisation Docker, reverse proxy Nginx avec support SSL/TLS, base de donnees PostgreSQL, cache Redis et stockage objet MinIO. Cette couche est orchestree par Docker Compose, qui gere le cycle de vie des cinq services conteneurises et assure la persistance des donnees via des volumes dedies.")

# 3.2.2 Architecture physique
add_heading(doc, '3.2.2 Architecture physique et de deploiement', 2)

add_para(doc, "L'architecture physique de deploiement du systeme GED-ISIPA repose sur une stack Docker Compose orchestreant cinq services conteneurises. Cette architecture a ete concue pour garantir la portabilite, la reproductibilite et la facilite de deploiement, conformement aux bonnes pratiques de DevOps recommandees par la documentation officielle Docker (2024). L'objectif est de permettre le deploiement de l'application sur un serveur VPS (Virtual Private Server) avec un minimum de configuration manuelle.")

add_figure(doc, f'{DIAG_DIR}/fig_deployment.png', 'Figure 3.2 : Architecture de deploiement Docker Compose')

add_para(doc, "Le service applicatif principal est l'application Next.js, construite selon un processus de build multi-etapes (multi-stage build) produisant une image Docker optimisee. L'etape de build installe les dependances, genere le client Prisma, compile l'application TypeScript et produit un bundle optimise pour la production. L'etape de runtime utilise une image Node.js Alpine legere, copie uniquement les artefacts necessaires et expose le port 3000. Le processus de seeding de la base de donnees est execute automatiquement lors du demarrage, garantissant que les donnees initiales (organisations, utilisateurs de test, workflows par defaut) sont disponibles.")

add_para(doc, "Le service de base de donnees PostgreSQL assure le stockage persistant des donnees applicatives en production. Configure avec un utilisateur dedie (aeip) et une base de donnees specifique (aeip_enterprise), il expose le port standard 5432 et utilise un volume Docker pour la persistance des donnees. Le service Redis fournit un cache en memoire pour les sessions et les donnees frequemment consultees, ameliorant les performances de l'application en reduisant le nombre de requetes a la base de donnees. Le service MinIO offre un stockage d'objets compatible S3 pour les fichiers documentaires, avec une console d'administration accessible sur le port 9001.")

add_para(doc, "Le service Nginx agit comme reverse proxy, terminant les connexions HTTPS et redistribuant les requetes vers l'application Next.js. Sa configuration inclut la compression gzip, la gestion des connexions WebSocket pour le hot reload en developpement, et le support SSL/TLS avec des certificats auto-signes ou letsencrypt. Les ports 80 et 443 sont exposes publiquement, tandis que les services internes communiquent via le reseau Docker dedie.")

# 3.2.3 Architecture applicative
add_heading(doc, '3.2.3 Architecture applicative', 2)

add_para(doc, "L'architecture applicative du systeme GED-ISIPA suit le patron d'architecture de l'App Router Next.js, qui organise le code selon une structure basee sur le systeme de fichiers. Chaque repertoire sous le chemin app/ correspond a une route de l'application, et les fichiers speciaux (page.tsx, layout.tsx, loading.tsx, error.tsx) definissent le comportement de chaque route. Cette approche convention-over-configuration permet une navigation intuitive dans le code source et facilite la maintenance.")

add_figure(doc, f'{DIAG_DIR}/fig_archi_applicative.png', 'Figure 3.3 : Architecture applicative Next.js App Router')

add_para(doc, "L'application est organisee en trois groupes de routes principaux : le groupe (auth) qui regroupe les pages d'authentification (connexion et inscription), le groupe (dashboard) qui contient l'espace de travail des utilisateurs d'une organisation (documents, archives, audit, modules, workflows, notifications, parametres), et le groupe admin qui rassemble les fonctions de super-administration de la plateforme (organisations, facturation, analyses, modules globaux). Chaque groupe possede son propre layout, assurant la coherence visuelle et fonctionnelle au sein du groupe.")

add_para(doc, "Les bibliotheques metier sont centralisees dans le repertoire src/lib/, avec des modules dedies pour la configuration d'authentification (auth.ts), la matrice de permissions (permissions.ts), le moteur de workflow (workflow.ts), le moteur de modules (module-engine.ts), le generateur de jetons d'organisation (token-engine.ts), le moteur de redirection (redirection.ts), les constantes de l'application (constants.ts), les validateurs Zod (validators.ts) et l'acces a la base de donnees (db.ts). Cette separation entre la logique metier et la logique de presentation respecte le principe de responsabilite unique (SRP) tel que defini par Martin (2017).")

# 3.2.4 Flux de donnees
add_heading(doc, '3.2.4 Flux de donnees', 2)

add_para(doc, "Le flux de donnees typique dans le systeme GED-ISIPA suit le pattern suivant : lorsqu'un utilisateur effectue une action sur l'interface, la requete HTTP est d'abord interceptee par le middleware Next.js, qui verifie l'authenticite du jeton JWT, applique le rate limiting, injecte les en-tetes de securite et verifie les droits d'acces grossiers (par exemple, seuls les SUPER_ADMIN peuvent acceder aux routes /admin/*). Si la requete est autorisee, elle est transmise au Route Handler correspondant, qui effectue une verification RBAC fine-grained, valide les donnees entrantes avec Zod, execute la logique metier via Prisma ORM et retourne une reponse JSON.")

add_figure(doc, f'{DIAG_DIR}/fig_data_flow.png', 'Figure 3.4 : Diagramme de flux de donnees du systeme GED-ISIPA')

add_para(doc, "Pour les operations de workflow, le flux est legerement different : l'utilisateur declenche une transition de workflow via l'interface, la requete est transmise a l'endpoint /api/workflows/[id]/execute, qui charge la transition avec ses etats source et cible, verifie que l'utilisateur possede un role autorise dans la liste allowedRoles de la transition, appelle la fonction canTransition() du moteur de workflow pour valider la coherence de la transition, met a jour l'etat du workflow et synchronise automatiquement le statut du document associe. Ce flux garantit que chaque transition est effectuee dans le respect des regles metier et de la tracabilite.")

# ============ 3.3 TECHNOLOGIES UTILISEES ============
add_heading(doc, '3.3 Technologies utilisees', 1)

add_heading(doc, '3.3.1 Stack technologique principal', 2)

add_para(doc, "Le choix des technologies utilisees pour le developpement de GED-ISIPA a ete guide par plusieurs criteres : la productivite du developpeur, la performance de l'application, la disponibilite de la communaute et de la documentation, la compatibilite avec les requirements du projet, et l'adéquation avec les competences de l'equipe de developpement. Le tableau suivant presente l'ensemble du stack technologique retenu.")

add_table(doc,
    ['Composant', 'Technologie', 'Version', 'Role'],
    [
        ['Framework', 'Next.js', '16.1.1', 'Framework full-stack React avec SSR/SSG'],
        ['Bibliotheque UI', 'React', '19.x', 'Bibliotheque de composants declaratifs'],
        ['Langage', 'TypeScript', '5.x', 'Typage statique et securite au compile-time'],
        ['ORM', 'Prisma', '6.11.1', 'Mapping objet-relationnel type-safe'],
        ['Base de donnees (dev)', 'SQLite', '3.x', 'Base embarquee pour le developpement'],
        ['Base de donnees (prod)', 'PostgreSQL', '16', 'Base relationnelle pour la production'],
        ['Authentification', 'NextAuth.js', '4.24.11', 'Authentification JWT et gestion de sessions'],
        ['Composants UI', 'shadcn/ui + Radix', 'Derniere', 'Composants accessibles et personnalisables'],
        ['CSS', 'Tailwind CSS', '4.x', 'Framework CSS utility-first'],
        ['Validation', 'Zod', '3.x', 'Validation de schemas TypeScript'],
        ['Conteneurisation', 'Docker + Compose', '24.x', 'Deploiement conteneurise'],
        ['Reverse Proxy', 'Nginx', '1.25', 'Proxy inverse avec SSL et compression'],
        ['Cache', 'Redis', '7.x', 'Cache en memoire pour les sessions'],
        ['Stockage fichiers', 'MinIO', 'latest', 'Stockage d\'objets compatible S3'],
    ],
    [3.5, 3.5, 2, 8]
)

add_para(doc, "Tableau 3.1 : Stack technologique de GED-ISIPA", bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

add_heading(doc, '3.3.2 Justification des choix', 2)

add_para(doc, "Le choix de Next.js comme framework principal se justifie par plusieurs avantages determinants. Premierement, Next.js offre un mode de rendu hybride (Server-Side Rendering et Client-Side Rendering) permettant d'optimiser les performances de chargement des pages et le referencement, comme documente par Vercel (2024). Deuxiemement, l'App Router permet de structurer l'application selon une architecture basee sur le systeme de fichiers, facilitant l'organisation du code et la navigation. Troisiemement, Next.js integre nativement les Route Handlers, qui remplacent avantageusement les API Routes traditionnelles en offrant un meilleur typage et une meilleure integration avec le systeme de cache. Quatriemement, le framework beneficie d'un ecosysteme mature et d'une communaute active, avec des utilisateurs majeurs comme Netflix, TikTok et Notion.")

add_para(doc, "Le choix de Prisma comme ORM est justifie par son approche declarative du schema de donnees, sa generation automatique de types TypeScript a partir du schema, et son support natif de SQLite en developpement et PostgreSQL en production. Le schema Prisma sert egalement de documentation vivante du modele de donnees, eliminant la necessite de maintenir un schema SQL separe. Comme le souligne la documentation officielle Prisma (2024), l'approche schema-first garantit la coherence entre le modele de donnees et le code applicatif.")

add_para(doc, "Le choix de NextAuth.js repond a la necessite d'integrer rapidement un systeme d'authentification robuste dans l'ecosysteme Next.js. La strategie JWT a ete privilegiee par rapport aux sessions en base de donnees pour des raisons de performance et de scalabilite : chaque requete peut etre authentifiee sans acces a la base de donnees, reduisant la latence et la charge sur le serveur. Cette approche est conforme aux recommandations de Fielding (2000) sur les architectures RESTful, ou l'etat de session est delegue au client.")

add_para(doc, "L'adoption de shadcn/ui plutot que d'une bibliotheque de composants traditionnelle (comme Material UI ou Ant Design) se justifie par le fait que shadcn/ui ne constitue pas une dependance externe : les composants sont copies dans le projet et peuvent etre modifies librement. Cette approche offre un controle total sur le comportement et l'apparence des composants, tout en beneficiant de l'accessibilite native des primitives Radix UI.")

add_heading(doc, '3.3.3 Limites identifiees', 2)

add_para(doc, "Malgre les avantages presentes, certaines limites technologiques doivent etre mentionnees de maniere transparente. Premierement, l'authentification actuelle utilise une comparaison de mots de passe en texte clair, une approche acceptable uniquement en phase de developpement mais inacceptable en production. L'implementation du hachage bcrypt est planifiee comme priorite absolue. Deuxiemement, le rate limiting est implemente en memoire via une Map JavaScript, ce qui ne fonctionne pas en mode multi-instance. La migration vers Redis pour le stockage des compteurs de rate limiting est prevue. Troisiemement, les statistiques des tableaux de bord sont actuellement calculees a partir de donnees statiques (hardcodees) plutot qu'a partir de requetes en temps reel sur la base de donnees. Quatriemement, les notifications sont generees localement et non persistees dynamiquement. Ces limites sont explicitement reconnues et des solutions d'amelioration sont proposees dans la section 3.11.")

# ============ 3.4 MODELISATION DU SYSTEME ============
add_heading(doc, '3.4 Modelisation du systeme', 1)

add_para(doc, "La modelisation du systeme GED-ISIPA suit la methode Merise presentee au Chapitre I, en adoptant les trois niveaux de abstraction : conceptuel (MCD), logique (MLD) et physique (MPD). En complement, les diagrammes UML (cas d'utilisation, classes, sequence, activite, composants, deploiement) sont utilises pour decrire les aspects comportementaux et structurels du systeme, conformement aux recommandations de Booch, Rumbaugh et Jacobson (2005). Cette approche combinee Merise-UML permet de couvrir l'ensemble des dimensions du systeme, des donnees aux traitements, de la structure au comportement.")

add_heading(doc, '3.4.1 Diagramme de cas d\'utilisation', 2)

add_para(doc, "Le diagramme de cas d'utilisation identifie les interactions principales entre les acteurs du systeme et les fonctionnalites offertes. Quatre acteurs principaux ont ete identifies : le Super Administrateur (SUPER_ADMIN), responsable de la gestion globale de la plateforme ; l'Administrateur d'organisation (ORG_ADMIN), gestionnaire des utilisateurs et des processus au sein de son organisation ; le Createur (PROFESSOR, USER), utilisateur creeant et soumettant des documents ; et le Lecteur (VIEWER), utilisateur disposant d'un acces en lecture seule. Les cas d'utilisation sont organises en cinq packages fonctionnels : Authentification, Gestion documentaire, Workflow, Administration et Super Administration.")

add_figure(doc, f'{DIAG_DIR}/fig_use_case.png', 'Figure 3.5 : Diagramme de cas d\'utilisation du systeme GED-ISIPA')

add_para(doc, "Ce diagramme illustre la couverture fonctionnelle du systeme et la repartition des responsabilites entre les differents profils d'utilisateurs. On constate que le Createur dispose des droits de creation, modification, soumission et partage de documents, tandis que l'Administrateur d'organisation dispose des droits d'approbation, de publication et de gestion des utilisateurs. Le Super Administrateur, quant a lui, a acces aux fonctions de gestion multi-organisation et de supervision de la plateforme. Cette separation des roles est essentielle pour garantir la securite et la tracabilite des operations, conformement aux principes RBAC (Role-Based Access Control) definis par Ferraiolo et al. (2001).")

add_heading(doc, '3.4.2 Modele Conceptuel de Donnees (MCD)', 2)

add_para(doc, "Le modele conceptuel de donnees de GED-ISIPA comprend quatorze entites principales, reliees par des associations reflechissant les regles de gestion du systeme. Ce modele a ete elabore en suivant la methode MERISE presentee au Chapitre I, avec une approche d'abstraction progressive des donnees telle que decrite par Tardieu, Rochfeld et Colletti (1991). L'entite Organization constitue la racine de l'architecture multi-tenant : chaque organisation possede un identifiant unique, un nom, un slug pour les URLs, un code AEIP genere automatiquement et un type parmi huit possibilites.")

add_figure(doc, f'{DIAG_DIR}/fig_mcd.png', 'Figure 3.6 : Modele Conceptuel de Donnees (MCD)')

add_para(doc, "L'entite Organization est l'entite centrale de l'architecture multi-tenant. Chaque organisation possede un identifiant unique (id), un nom, un slug unique pour les URLs, un code AEIP genere automatiquement (par exemple AEIP-UNI-SXKLQT pour l'ISIPA), un type parmi huit possibilites (UNIVERSITY, HOSPITAL, COMPANY, GOVERNMENT, SME, INSTITUTION, NGO, LAW_FIRM), un statut (ACTIVE, SUSPENDED, TRIAL, CHURNED), et des informations de personnalisation (logo, couleur primaire, description). La cardinalite entre Organization et User est de type 1,N, signifiant qu'une organisation contient plusieurs utilisateurs, et chaque utilisateur appartient a exactement une organisation.")

add_para(doc, "L'entite Document constitue le coeur fonctionnel de l'application. Un document est defini par son titre, sa reference unique generee automatiquement, son type parmi quatorze categories, son statut dans le cycle de vie (DRAFT, PENDING_REVIEW, APPROVED, PUBLISHED, ARCHIVED, REJECTED), et son niveau de classification (PUBLIC, INTERNAL, CONFIDENTIAL, RESTRICTED). Chaque document est rattache a une organisation (isolation multi-tenant), un auteur (User), un departement optionnel et un workflow de validation. Les cardinalites montrent qu'un utilisateur peut creer plusieurs documents (1,N), qu'un departement peut contenir plusieurs documents (0,N) et qu'un document possede un historique de versions (1,N).")

add_heading(doc, '3.4.3 Modele Logique de Donnees (MLD)', 2)

add_para(doc, "Le modele logique de donnees constitue la traduction du MCD en termes de tables relationnelles, conformement a la demarche Merise. Chaque entite du MCD correspond a une table, et chaque association de type un-a-plusieurs est traduite par l'ajout d'une cle etrangere dans la table du cote plusieurs. Le MLD introduit egalement les types de donnees, les contraintes d'integrite (cles primaires, cles etrangeres, contraintes d'unicite) et les index necessaires aux performances des requetes.")

add_figure(doc, f'{DIAG_DIR}/fig_mld.png', 'Figure 3.7 : Modele Logique de Donnees (MLD)')

add_para(doc, "La table User illustre la traduction des associations du MCD : la cle etrangere organizationId reference la table Organization (traduction de l'appartenance d'un utilisateur a une organisation), et la cle etrangere departmentId reference la table Department (traduction de l'affectation optionnelle a un departement). La contrainte d'unicite sur le couple (email, organizationId) garantit que chaque adresse email est unique au sein d'une organisation, mais peut etre reutilisee dans differentes organisations, ce qui est conforme a la logique multi-tenant du systeme.")

add_heading(doc, '3.4.4 Modele Physique de Donnees (MPD)', 2)

add_para(doc, "Le modele physique de donnees represente l'implementation concrete des tables relationnelles dans le systeme de gestion de base de donnees. Il precise les types de donnees physiques (VARCHAR, TEXT, INTEGER, DATETIME, BOOLEAN), les contraintes NOT NULL, DEFAULT et CHECK, ainsi que les index optimises pour les requetes les plus frequentes. Le MPD constitue le lien direct entre le modele logique et le schema Prisma effectivement implemente.")

add_figure(doc, f'{DIAG_DIR}/fig_mpd.png', 'Figure 3.8 : Modele Physique de Donnees (MPD)')

add_para(doc, "Le MPD revele plusieurs choix d'implementation significatifs. Premierement, tous les identifiants utilisent le format CUID (Collision-resistant Unique Identifiers), genere automatiquement par Prisma, garantissant l'unicite globale meme dans un contexte distribue. Deuxiemement, le champ allowedRoles de la table WorkflowTransition est de type JSON, permettant de stocker un tableau de roles autorises pour chaque transition sans necessiter une table de liaison supplementaire. Troisiemement, chaque table contenant des donnees organisationnelles inclut un champ organizationId, indexe pour optimiser les requetes de filtrage multi-tenant. Quatriemement, les champs de timestamp (createdAt, updatedAt) sont generes automatiquement par Prisma, assurant la tracabilite temporelle de toutes les entites.")

add_heading(doc, '3.4.5 Dictionnaire de donnees', 2)

add_para(doc, "Le dictionnaire de donnees presente de maniere exhaustive l'ensemble des proprietes de l'entite Document, entite centrale du systeme. Ce dictionnaire constitue un outil de reference essentiel pour la comprehension et la maintenance du systeme.")

add_table(doc,
    ['Champ', 'Type', 'Contraintes', 'Description'],
    [
        ['id', 'String (CUID)', 'PK, auto', 'Identifiant unique du document'],
        ['title', 'String', 'NOT NULL', 'Titre du document'],
        ['reference', 'String', 'UNIQUE', 'Reference unique auto-generee'],
        ['type', 'DocumentType', 'NOT NULL', 'Type parmi 14 categories'],
        ['status', 'DocumentStatus', 'NOT NULL, DEFAULT DRAFT', 'Statut dans le cycle de vie'],
        ['classification', 'Classification', 'NOT NULL, DEFAULT INTERNAL', 'Niveau de classification'],
        ['content', 'String', 'OPTIONAL', 'Contenu textuel du document'],
        ['fileUrl', 'String', 'OPTIONAL', 'URL du fichier physique'],
        ['fileName', 'String', 'OPTIONAL', 'Nom du fichier original'],
        ['fileSize', 'Int', 'OPTIONAL', 'Taille du fichier en octets'],
        ['mimeType', 'String', 'OPTIONAL', 'Type MIME du fichier'],
        ['organizationId', 'String', 'FK Organization, NOT NULL, INDEX', 'Cle etrangere multi-tenant'],
        ['authorId', 'String', 'FK User, NOT NULL', 'Cle etrangere vers l\'auteur'],
        ['departmentId', 'String', 'FK Department, OPTIONAL', 'Departement rattache'],
        ['workflowId', 'String', 'FK Workflow, OPTIONAL', 'Workflow de validation'],
        ['currentWorkflowStateId', 'String', 'FK WorkflowState, OPTIONAL', 'Etat courant du workflow'],
        ['archivedAt', 'DateTime', 'OPTIONAL', 'Date d\'archivage'],
        ['archivedBy', 'String', 'FK User, OPTIONAL', 'Utilisateur ayant archive'],
        ['version', 'Int', 'NOT NULL, DEFAULT 1', 'Numero de version courant'],
        ['createdAt', 'DateTime', 'AUTO', 'Date de creation'],
        ['updatedAt', 'DateTime', 'AUTO', 'Date de derniere modification'],
    ],
    [3.5, 3, 4, 6]
)
add_para(doc, "Tableau 3.2 : Dictionnaire de donnees - Entite Document", bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

add_heading(doc, '3.4.6 Enumerations du systeme', 2)

add_para(doc, "Le systeme GED-ISIPA definit dix enumerations au niveau du schema Prisma, chacune representant un ensemble fini de valeurs autorisees pour un champ donne. Ces enumerations assurent la coherence des donnees et la validation au niveau de la base de donnees.")

add_table(doc,
    ['Enumeration', 'Nombre de valeurs', 'Valeurs principales'],
    [
        ['OrganizationType', '8', 'UNIVERSITY, HOSPITAL, COMPANY, GOVERNMENT, SME, INSTITUTION, NGO, LAW_FIRM'],
        ['Role', '14', 'SUPER_ADMIN, ORG_ADMIN, MANAGER, USER, VIEWER, DEAN, PROFESSOR, DOCTOR, NURSE, LAWYER, PARALEGAL, CFO, HR_MANAGER, CIVIL_SERVANT'],
        ['DocumentType', '14', 'ACADEMIC_RECORD, ADMINISTRATIVE, FINANCIAL, CORRESPONDENCE, REPORT, CONTRACT, CERTIFICATE, MEMO, POLICY, MEDICAL_RECORD, LEGAL_BRIEF, INVOICE, PROPOSAL, OTHER'],
        ['DocumentStatus', '6', 'DRAFT, PENDING_REVIEW, APPROVED, PUBLISHED, ARCHIVED, REJECTED'],
        ['Classification', '4', 'PUBLIC, INTERNAL, CONFIDENTIAL, RESTRICTED'],
        ['SubscriptionPlan', '4', 'FREE, STARTER, PROFESSIONAL, ENTERPRISE'],
        ['OrganizationStatus', '4', 'ACTIVE, SUSPENDED, TRIAL, CHURNED'],
        ['ModuleStatus', '3', 'ACTIVE, SUSPENDED, INACTIVE'],
        ['SubscriptionStatus', '4', 'ACTIVE, PAST_DUE, CANCELED, TRIAL'],
        ['AuditAction', '17', 'CREATE, READ, UPDATE, DELETE, ARCHIVE, RESTORE, DOWNLOAD, SHARE, APPROVE, REJECT, LOGIN, LOGOUT, MODULE_ACTIVATE, MODULE_SUSPEND, WORKFLOW_EXECUTE, ORGANIZATION_CREATE, ORGANIZATION_SUSPEND'],
    ],
    [3.5, 2.5, 11]
)
add_para(doc, "Tableau 3.3 : Enumerations du schema Prisma", bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

add_heading(doc, '3.4.7 Diagramme de classes', 2)

add_para(doc, "Le diagramme de classes presente la structure orientee objet du systeme, montrant les 14 modeles de donnees avec leurs attributs, methodes et relations. Ce diagramme constitue la traduction UML du schema Prisma et permet de visualiser l'architecture des donnees sous un angle complementaire du MCD Merise. Conformement aux recommandations de Rumbaugh, Jacobson et Booch (2004), le diagramme de classes est le artefact central de la modelisation orientee objet.")

add_figure(doc, f'{DIAG_DIR}/fig_class_diagram.png', 'Figure 3.9 : Diagramme de classes du systeme GED-ISIPA')

add_para(doc, "Le diagramme de classes met en evidence la centralite du modele Document, qui entretient des relations avec la majorite des autres modeles. On observe egalement le motif de composition entre Workflow, WorkflowState et WorkflowTransition, formant le moteur de workflow du systeme. La relation d'agregation entre Organization et OrganizationModule illustre le systeme de modules activables par organisation. Les cardinalites sont exprimees en notation UML standard (1..*, 0..*, 1..1), permettant une lecture immediate des contraintes relationnelles.")

# Save intermediate progress
doc.save(OUTPUT)
print(f"Part 1 saved: {os.path.getsize(OUTPUT)/1024:.1f} KB")
