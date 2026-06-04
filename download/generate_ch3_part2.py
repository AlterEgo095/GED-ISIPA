#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Chapitre 3 GED-ISIPA - Part 2: Sections 3.5-3.7 (Implementation, Features, Interfaces with screenshots)"""
import os, sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT = '/home/z/my-project/download/Chapitre3_GED_ISIPA_Definitif.docx'
DIAG_DIR = '/home/z/my-project/download/ch3_diagrams'
SS_DIR = '/home/z/my-project/download/screenshots'

# ============ Helper Functions (same as Part 1) ============
def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F) if level <= 2 else RGBColor(0x2C, 0x3E, 0x50)
    return h

def add_para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=11, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    return p

def add_figure(doc, img_path, caption, width_cm=15):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.name = 'Calibri'
    r.bold = True
    r.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '1E3A5F')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9)
                r.font.name = 'Calibri'
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = 'Calibri'
            if ri % 2 == 1:
                set_cell_shading(cell, 'F0F4F8')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    for r in p.runs:
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
    return p

# Load existing document
doc = Document(OUTPUT)

# ============ 3.5 REALISATION ET IMPLEMENTATION ============
add_heading(doc, '3.5 Realisation et implementation', 1)

add_heading(doc, '3.5.1 Structure du projet', 2)

add_para(doc, "Le projet GED-ISIPA est organise selon la structure standard d'une application Next.js 16 avec App Router. Le repertoire source (src/) contient les sous-repertoires app/ pour les routes et pages, components/ pour les composants React, lib/ pour les modules metier, et le repertoire prisma/ pour le schema de base de donnees et les scripts de seeding. L'organisation du code suit le principe de proximite fonctionnelle : les fichiers lies a une meme fonctionnalite sont regroupes dans un meme repertoire. Par exemple, les composants d'interface lies aux documents sont dans components/documents/, les composants d'authentification dans components/auth/, et les composants de mise en page dans components/layout/.")

add_para(doc, "Le fichier le plus critique du projet est le schema Prisma (prisma/schema.prisma), qui definit l'ensemble du modele de donnees et sert de contrat entre la base de donnees et le code applicatif. Ce schema est utilise par Prisma pour generer automatiquement le client TypeScript (prisma generate), garantissant que le typage des donnees dans le code est toujours synchronise avec la structure reelle de la base de donnees. Le fichier de seed (prisma/seed.ts) initialise la base de donnees avec trois organisations de test (ISIPA, Hopital Central, Ministere du Plan), dix utilisateurs couvrant sept roles differents, et 21 documents representatifs.")

add_heading(doc, '3.5.2 Gestion des roles et permissions', 2)

add_para(doc, "Le systeme de gestion des roles et permissions de GED-ISIPA constitue l'un des aspects les plus elabores de l'application. Il repose sur une matrice de permissions statique definie dans le fichier permissions.ts, qui croise quatorze roles avec neuf ressources et douze types d'actions. Cette approche RBAC (Role-Based Access Control) est conforme au standard propose par Ferraiolo et al. (2001) et aux recommandations du NIST (2020) en matiere de controle d'acces.")

add_figure(doc, f'{DIAG_DIR}/fig_archi_logique.png', 'Figure 3.10 : Hierarchie des roles RBAC avec niveaux d\'autorite', width_cm=14)

add_para(doc, "La hierarchie des roles est organisee selon un systeme de niveaux d'autorite allant de 100 (SUPER_ADMIN) a 10 (VIEWER). Chaque niveau herite implicitement des permissions des niveaux inferieurs, garantissant une coherence dans la distribution des droits. Le SUPER_ADMIN dispose d'un acces complet a toutes les ressources et actions, incluant la gestion multi-organisation. L'ORG_ADMIN gere les utilisateurs et les processus au sein de son organisation. Les roles metier (DEAN, PROFESSOR, DOCTOR, NURSE, LAWYER, PARALEGAL, CFO, HR_MANAGER, CIVIL_SERVANT) disposent de permissions adaptees a leur contexte professionnel. Enfin, les roles generiques (MANAGER, USER, VIEWER) offrent des niveaux d'acces progressifs.")

add_table(doc,
    ['Role', 'Niveau', 'Actions cles sur documents', 'Gestion utilisateurs'],
    [
        ['SUPER_ADMIN', '100', 'Toutes (CRUD + approbation + publication + archivage)', 'Plein controle + organisations'],
        ['ORG_ADMIN', '80', 'CRUD + approbation + publication + archivage', 'CRUD utilisateurs + departements'],
        ['DEAN', '70', 'CRUD + approbation + publication + archivage', 'Lecture + gestion'],
        ['CFO / HR_MANAGER', '65', 'CRUD + approbation + archivage', 'Lecture / Creation + gestion'],
        ['CIVIL_SERVANT / DOCTOR / LAWYER', '60', 'CRUD + approbation + archivage', 'Lecture uniquement'],
        ['MANAGER', '50', 'CRUD + approbation + rejet + archivage', 'Lecture uniquement'],
        ['PROFESSOR / NURSE / PARALEGAL', '30-40', 'CRUD + partage', 'Lecture uniquement'],
        ['USER', '20', 'Create + read + update + share', 'Lecture uniquement'],
        ['VIEWER', '10', 'Lecture uniquement', 'Lecture uniquement'],
    ],
    [4, 1.5, 5.5, 5.5]
)
add_para(doc, "Tableau 3.4 : Synthese des permissions par role", bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

add_heading(doc, '3.5.3 Moteur de workflow', 2)

add_para(doc, "Le moteur de workflow de GED-ISIPA est un systeme configurable de machines a etats finis permettant de modeliser les circuits de validation documentaire. Chaque organisation dispose d'au moins un workflow, cree automatiquement lors du seed de la base de donnees. Le workflow par defaut comporte cinq etats (Brouillon, En revision, Approuve, Publie, Rejete) et cinq transitions regissant le passage d'un etat a l'autre.")

add_para(doc, "L'implementation du moteur de workflow repose sur trois modeles Prisma : Workflow, WorkflowState et WorkflowTransition. Le workflow est cree a partir d'une configuration typee (WorkflowConfig) via la fonction createWorkflowFromConfig, qui cree le workflow, ses etats et ses transitions en une seule transaction Prisma. Chaque transition definit un etat source, un etat cible, un nom et une liste de roles autorises (allowedRoles). La fonction canTransition() verifie que la transition demandee est valide en verifiant que l'etat courant correspond a l'etat source de la transition et que l'utilisateur possede un role autorise.")

add_para(doc, "Le systeme de transitions est extensible : chaque organisation peut definir des workflows personnalises via l'interface de creation de workflows, avec des etats et des transitions adaptes a ses besoins specifiques. Par exemple, le workflow medical d'un hopital pourrait inclure des etats supplementaires comme 'En cours d\'examen' ou 'En attente de signature', tandis que le workflow juridique d'un cabinet pourrait inclure des etats de 'Due diligence' ou 'En negociation'. Le couplage entre le moteur de workflow et le statut du document est automatique : chaque transition de workflow met a jour le champ status du document, garantissant la coherence entre l'etat du workflow et le statut fonctionnel du document.")

add_heading(doc, '3.5.4 Moteur de modules', 2)

add_para(doc, "Le systeme de modules de GED-ISIPA permet a chaque organisation d'activer ou de desactiver des fonctionnalites selon ses besoins specifiques. Quinze modules sont definis dans le catalogue, chacun associe a un ou plusieurs types d'organisations et potentiellement a des dependances envers d'autres modules. Le moteur de modules gere automatiquement les dependances entre modules : par exemple, le module Bibliotheque depend du module Academique, et ne peut etre active que si ce dernier est deja actif.")

add_figure(doc, f'{DIAG_DIR}/fig_archi_applicative.png', 'Figure 3.11 : Systeme de modules par type d\'organisation', width_cm=14)

add_para(doc, "Les 15 modules du catalogue couvrent l'ensemble des besoins fonctionnels identifiés pour les huit types d'organisations supportees. Pour les universites : RH (commun), Academique, Bibliotheque (depend d'Academique), Recherche (depend d'Academique). Pour les hopitaux : Medical, Pharmacie (depend de Medical), Laboratoire (depend de Medical). Pour les entreprises : Finance, CRM, Projets. Pour les gouvernements : Procedures, Archivage, Conformite. Pour les cabinets juridiques : Juridique, Facturation (depend de Juridique). Le module RH est le seul commun a tous les types d'organisations, refletant le caractere transversal de la gestion des ressources humaines.")

add_heading(doc, '3.5.5 Authentification et securite', 2)

add_para(doc, "Le systeme d'authentification de GED-ISIPA est base sur NextAuth.js en strategie JWT. Lorsqu'un utilisateur soumet ses identifiants (email et mot de passe, avec un code d'organisation optionnel), le fournisseur d'identite CredentialsProvider effectue les verifications suivantes : recherche de l'utilisateur par email avec inclusion des informations d'organisation, verification que le compte est actif (isActive), comparaison du mot de passe, verification optionnelle du code d'organisation, mise a jour de la date de derniere connexion (lastLogin), et creation d'une entree dans le journal d'audit (action LOGIN). Le jeton JWT genere contient les claims suivants : id, role, organizationId, organizationName, organizationType, organizationCode et departmentId.")

add_para(doc, "Le middleware Next.js (middleware.ts) constitue la premiere ligne de defense de l'application. Il intercepte chaque requete entrante et applique les mesures de securite suivantes : injection d'en-tetes de securite HTTP (X-Frame-Options: DENY, X-Content-Type-Options: nosniff, Referrer-Policy: strict-origin-when-cross-origin, X-XSS-Protection: 1; mode=block), rate limiting (30 requetes par minute pour les endpoints d'authentification, 100 pour les API generales), protection des routes sensibles (redirection des utilisateurs non authentifies vers la page de connexion), et verification du role SUPER_ADMIN pour les routes d'administration plateforme.")

add_para(doc, "L'isolation multi-tenant est assuree par le filtre systematique de toutes les requetes de base de donnees sur la base de l'identifiant d'organisation extrait du jeton JWT. Chaque endpoint API recupere l'organizationId du jeton et l'utilise comme filtre dans toutes les requetes Prisma, garantissant qu'un utilisateur ne peut jamais acceder aux donnees d'une autre organisation. Cette approche d'isolation applicative est appropriee pour le contexte actuel et pourrait etre renforcee par une isolation au niveau du schema PostgreSQL dans une version future, comme recommande par les bonnes pratiques de multi-tenancy.")

add_heading(doc, '3.5.6 API RESTful', 2)

add_para(doc, "L'API de GED-ISIPA est exposee via 28 fichiers de Route Handlers, couvrant 41 endpoints distincts. Chaque endpoint suit le pattern REST standard : GET pour la lecture, POST pour la creation, PUT pour la mise a jour, DELETE pour la suppression. Tous les endpoints sont proteges par une authentification JWT (sauf /api/health et /api/auth), un controle RBAC fine-grained et une validation des donnees entrantes avec Zod.")

add_table(doc,
    ['Endpoint', 'Methodes', 'Description', 'Permission requise'],
    [
        ['/api/documents', 'GET, POST', 'Liste paginee + creation', 'read / create (documents)'],
        ['/api/documents/[id]', 'GET, PUT, DELETE', 'CRUD document individuel', 'read / update / delete'],
        ['/api/documents/[id]/submit', 'POST', 'Soumettre en revision', 'create (documents)'],
        ['/api/documents/[id]/approve', 'POST', 'Approuver un document', 'approve (documents)'],
        ['/api/documents/[id]/publish', 'POST', 'Publier un document', 'publish (documents)'],
        ['/api/documents/[id]/archive', 'POST', 'Archiver un document', 'archive (documents)'],
        ['/api/documents/[id]/restore', 'POST', 'Restaurer depuis archive', 'restore (documents)'],
        ['/api/documents/[id]/versions', 'GET', 'Historique des versions', 'read (documents)'],
        ['/api/users', 'GET, POST', 'Liste + creation utilisateurs', 'read / create (users)'],
        ['/api/organizations', 'GET, POST', 'Liste + creation organisations', 'SUPER_ADMIN uniquement'],
        ['/api/workflows', 'GET, POST', 'Liste + creation workflows', 'read / create (workflows)'],
        ['/api/workflows/[id]/execute', 'POST', 'Executer une transition', 'role dans allowedRoles'],
        ['/api/audit', 'GET', 'Journal d\'audit filtre', 'read (audit)'],
        ['/api/dashboard', 'GET', 'Statistiques du tableau de bord', 'Authentification requise'],
        ['/api/search', 'GET', 'Recherche globale', 'Authentification requise'],
        ['/api/health', 'GET', 'Verification de sante', 'Aucune'],
    ],
    [4.5, 2, 4.5, 4.5]
)
add_para(doc, "Tableau 3.5 : Endpoints API principaux", bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

add_heading(doc, '3.5.7 Systeme de jetons d\'organisation', 2)

add_para(doc, "Chaque organisation recoit un code d'identification unique genere automatiquement au format AEIP-PREFIX-CODE, ou PREFIX est un code de trois lettres correspondant au type d'organisation et CODE est une chaine alphanumerique de six caracteres generée aleatoirement. Ce systeme de jetons permet d'identifier de maniere unique chaque organisation et de verifier l'appartenance d'un utilisateur lors de l'authentification.")

add_table(doc,
    ['Type d\'organisation', 'Prefixe', 'Exemple'],
    [
        ['Universite', 'UNI', 'AEIP-UNI-SXKLQT'],
        ['Hopital', 'HOS', 'AEIP-HOS-LZYNWA'],
        ['Entreprise', 'COR', 'AEIP-COR-X9M2KP'],
        ['Gouvernement', 'GOV', 'AEIP-GOV-ELHBGX'],
        ['PME', 'PME', 'AEIP-PME-R4T7YN'],
        ['Institution', 'INS', 'AEIP-INS-W8Q3DL'],
        ['ONG', 'ONG', 'AEIP-ONG-F5V9HZ'],
        ['Cabinet juridique', 'JUR', 'AEIP-JUR-B6N2XK'],
    ],
    [5, 3, 5]
)
add_para(doc, "Tableau 3.6 : Correspondance des prefixes de jetons AEIP", bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

# ============ 3.6 FONCTIONNALITES ============
add_heading(doc, '3.6 Presentation detaillee des fonctionnalites', 1)

add_heading(doc, '3.6.1 Gestion documentaire', 2)

add_para(doc, "La gestion documentaire constitue le coeur fonctionnel de la plateforme GED-ISIPA. Elle permet aux utilisateurs de creer, consulter, modifier, approuver, publier et archiver des documents electroniques. Chaque document est caracterise par un titre, une reference unique generee automatiquement, un type parmi quatorze categories, un statut dans le cycle de vie, et un niveau de classification. La plateforme implemente les principes de l'archivage numerique inspirés du modele OAIS (Open Archival Information System) defini par la norme ISO 14721:2012. Dans ce cadre, le processus de soumission d'un document correspond au SIP (Submission Information Package), le document approuve et publie constitue l'AIP (Archival Information Package), et la restitution d'un document archive correspond au DIP (Dissemination Information Package).")

add_para(doc, "La recherche documentaire est disponible via un endpoint dedie (/api/search) permettant de rechercher par titre, reference, type, statut et classification. La pagination est implementee sur toutes les listes de documents, avec des limites par defaut et maximales pour eviter les surcharges. La gestion des versions est assuree par le modele DocumentVersion, qui conserve l'historique de chaque modification apportee a un document. Enfin, le telechargement de documents genere une entree dans le journal d'audit et le journal d'acces, garantissant la tracabilite complete des consultations.")

add_heading(doc, '3.6.2 Workflow de validation', 2)

add_para(doc, "Le workflow de validation permet de definir et d'executer des circuits de validation documentaire adaptes aux besoins de chaque organisation. Le workflow par defaut comporte cinq etats et cinq transitions, mais le systeme est extensible et permet la creation de workflows personnalises. Le controle des transitions est assure par une liste de roles autorises (allowedRoles) definie pour chaque transition, garantissant que seuls les utilisateurs disposant du role approprié peuvent effectuer une action donnee. Par exemple, seuls les roles USER, PROFESSOR, NURSE et PARALEGAL peuvent soumettre un document en revision, tandis que seuls les roles ORG_ADMIN, DEAN, DOCTOR, LAWYER, CFO et CIVIL_SERVANT peuvent approuver un document.")

add_heading(doc, '3.6.3 Systeme multi-tenant', 2)

add_para(doc, "Le systeme multi-tenant constitue l'innovation architecturale majeure de la plateforme. Il permet a plusieurs organisations de types differents d'utiliser la meme instance de l'application, tout en beneficiant d'une isolation totale de leurs donnees. L'isolation est assuree au niveau applicatif par le filtre systematique de toutes les requetes Prisma sur la base du champ organizationId extrait du jeton JWT de l'utilisateur. Cette approche, qualifiee d'isolation par colonne discriminante (discriminator column), est la methode la plus repandue pour l'implementation du multi-tenancy en architecture SaaS, comme documente par les travaux de Fowler (2002) sur les patterns d'architecture d'entreprise.")

add_heading(doc, '3.6.4 Tableaux de bord adaptatifs', 2)

add_para(doc, "La plateforme offre six tableaux de bord specialises, chacun adapte a un type d'organisation specifique. Le tableau de bord universitaire affiche des indicateurs academiques, le tableau de bord hospitalier presente des statistiques medicales, et le tableau de bord gouvernemental met en avant les procedures administratives. Chaque tableau de bord est construit a partir de composants reutilisables (StatCard, ChartCard, RecentList, QuickActions) definis dans le fichier dashboard-widgets.tsx, garantissant la coherence visuelle tout en permettant une personnalisation par type d'organisation.")

add_heading(doc, '3.6.5 Journal d\'audit', 2)

add_para(doc, "Le journal d'audit enregistre de maniere exhaustive toutes les actions effectuees sur la plateforme, incluant la creation, lecture, modification, suppression, archivage, restauration, telechargement, partage, approbation et rejet de documents, ainsi que les connexions, deconnexions, activations et suspensions de modules, les executions de workflow et les operations sur les organisations. Dix-sept types d'actions sont traces dans le schema AuditAction. Chaque entree d'audit contient l'identifiant de l'utilisateur, l'action effectuee, le type et l'identifiant de l'entite concernee, l'adresse IP et le user-agent du client. Cette tracabilite complete est essentielle pour la conformite reglementaire et repond aux exigences de la norme ISO 15489:2016 relative a la gestion des documents d'archive.")

add_heading(doc, '3.6.6 Gestion des utilisateurs et des departements', 2)

add_para(doc, "Le module de gestion des utilisateurs permet aux administrateurs d'organisation de creer, modifier et supprimer des comptes utilisateurs, de leur attribuer des roles et de les affecter a des departements. La gestion des departements permet de creer une structure hierarchique au sein de chaque organisation, facilitant l'organisation et le filtrage des documents par service. Chaque departement est defini par un nom, un code unique au sein de l'organisation, et est rattache a une organisation. Les utilisateurs peuvent etre affectes a un departement optionnel, permettant une organisation flexible des equipes.")

add_heading(doc, '3.6.7 Systeme de notifications', 2)

add_para(doc, "Le systeme de notifications permet d'informer les utilisateurs des evenements importants survenant dans la plateforme : approbation ou rejet d'un document, activation d'un module, rappels de taches a effectuer. Les notifications sont associees a chaque utilisateur et possedent un statut de lecture (lu/non lu). L'interface de notifications presente la liste des messages recents avec la possibilite de les marquer comme lus. A ce stade de developpement, les notifications sont generees localement et ne sont pas encore persistees dynamiquement via les API, ce qui constitue une limite identifiee dans la section 3.10.")

add_heading(doc, '3.6.8 Systeme de modules', 2)

add_para(doc, "Le systeme de modules offre une architecture fonctionnelle modulaire permettant a chaque organisation d'activer uniquement les fonctionnalites dont elle a besoin. Les 15 modules du catalogue couvrent les domaines fonctionnels suivants : Ressources Humaines (commun a toutes les organisations), Academique et Bibliotheque (universites), Medical, Pharmacie et Laboratoire (hopitaux), Finance, CRM et Projets (entreprises), Procedures, Archivage et Conformite (gouvernements), Juridique et Facturation (cabinets juridiques), et Recherche (universites). Le moteur de modules gere les dependances entre modules (par exemple, Bibliotheque depend d'Academique) et les contraintes de type d'organisation, empechant l'activation d'un module incompatible avec le type de l'organisation.")

# ============ 3.7 INTERFACES UTILISATEUR ============
add_heading(doc, '3.7 Presentation des interfaces utilisateur', 1)

add_para(doc, "Les interfaces utilisateur de GED-ISIPA ont ete concues selon les principes de l'ergonomie web moderne, en privilegiant la clarte, la coherence et l'accessibilite. L'ensemble des ecrans utilise la bibliotheque shadcn/ui avec un theme clair/sombre base sur Tailwind CSS, offrant une experience visuelle professionnelle et coherente. Chaque interface est presentee ci-dessous avec une capture d'ecran reelle de l'application deployee, accompagnee d'une analyse detaillee de son fonctionnement et de sa valeur metier.")

add_heading(doc, '3.7.1 Page de connexion', 2)

add_para(doc, "La page de connexion constitue le point d'entree de l'application. Elle presente un formulaire centre avec les champs email et mot de passe, ainsi qu'un champ optionnel pour le code d'organisation. Le formulaire est encadre dans une carte avec le logo GED-ISIPA et un message de bienvenue. La validation des champs est effectuee cote client avant la soumission, et les erreurs d'authentification sont affichees de maniere claire et non technique. Apres une connexion reussie, l'utilisateur est redirige automatiquement vers le tableau de bord correspondant a son type d'organisation, ou vers la console de super-administration s'il a le role SUPER_ADMIN.")

add_figure(doc, f'{SS_DIR}/01-login-page.png', 'Figure 3.12 : Page de connexion de l\'application GED-ISIPA', width_cm=14)

add_para(doc, "L'interface de connexion illustre le principe de simplicite volontaire : seuls les champs necessaires sont presents, reduisant la charge cognitive de l'utilisateur. Le champ optionnel 'Code d'organisation' permet une authentification contextuelle, verifiant que l'utilisateur appartient bien a l'organisation specifiee. Cette approche est particulierement utile dans un contexte multi-tenant ou un meme email pourrait theoriquement exister dans differentes organisations.")

add_heading(doc, '3.7.2 Page d\'inscription', 2)

add_para(doc, "La page d'inscription permet a un nouvel utilisateur de creer un compte sur la plateforme. Le formulaire collecte les informations essentielles : nom complet, adresse email, mot de passe, code d'organisation. L'inscription est liee a une organisation existante via le code AEIP, garantissant que chaque nouvel utilisateur est automatiquement rattache a son contexte organisationnel.")

add_figure(doc, f'{SS_DIR}/12-register-page.png', 'Figure 3.13 : Page d\'inscription', width_cm=14)

add_heading(doc, '3.7.3 Tableaux de bord specialises', 2)

add_para(doc, "Chaque type d'organisation dispose d'un tableau de bord adapte a son contexte. Le tableau de bord universitaire affiche les indicateurs academiques : nombre total de documents, etudiants, cours et publications, avec des graphiques de distribution par type de document et par statut. Le tableau de bord administrateur general presente une vue synthetique avec les statistiques globales de l'organisation.")

add_figure(doc, f'{SS_DIR}/02-dashboard-admin.png', 'Figure 3.14 : Tableau de bord administrateur general', width_cm=14)

add_para(doc, "Le tableau de bord universitaire presente des indicateurs specifiques au contexte academique, incluant le nombre de documents par type (dossiers academiques, administratifs, financiers), le nombre d'utilisateurs par role (professeurs, etudiants, personnel administratif) et les documents recemment soumis ou approuves. Cette specialisation du tableau de bord en fonction du type d'organisation est une caracteristique distinctive de GED-ISIPA.")

add_figure(doc, f'{SS_DIR}/16-university-dashboard.png', 'Figure 3.15 : Tableau de bord universitaire specialise', width_cm=14)

add_para(doc, "Le tableau de bord du directeur offre une vue synthetique des indicateurs cles de l'organisation, permettant un suivi rapide de l'activite documentaire et une prise de decision eclairee. Les widgets de statistiques, les graphiques de distribution et les listes de documents recents constituent les principaux composants de ce tableau de bord.")

add_figure(doc, f'{SS_DIR}/09-director-dashboard.png', 'Figure 3.16 : Tableau de bord du directeur', width_cm=14)

add_heading(doc, '3.7.4 Gestion documentaire', 2)

add_para(doc, "L'interface de gestion documentaire presente une vue en tableau des documents avec des colonnes pour le titre, la reference, le type, le statut, la classification, l'auteur et le departement. Des filtres dynamiques permettent de rechercher par texte, type, statut et classification. Un bouton de creation permet d'ajouter un nouveau document via un formulaire modal. Chaque ligne du tableau offre des actions contextuelles selon le statut du document et les permissions de l'utilisateur : consulter, modifier, soumettre, approuver, publier, archiver ou supprimer.")

add_figure(doc, f'{SS_DIR}/03-documents-list.png', 'Figure 3.17 : Liste des documents avec filtres et pagination', width_cm=14)

add_para(doc, "La vue de detail d'un document affiche l'ensemble des informations du document, incluant son statut dans le cycle de vie, son historique de versions, et les actions disponibles selon le role de l'utilisateur. L'interface met clairement en evidence le statut courant du document et les transitions de workflow possibles, guidant l'utilisateur dans le processus de validation.")

add_figure(doc, f'{SS_DIR}/04-document-detail.png', 'Figure 3.18 : Detail d\'un document avec actions contextuelles', width_cm=14)

add_heading(doc, '3.7.5 Archives', 2)

add_para(doc, "L'interface des archives presente la liste des documents archives, accessible uniquement aux utilisateurs disposant de la permission 'archive'. Elle reprend la meme structure que la liste des documents actifs, mais filtre automatiquement les documents dont le statut est ARCHIVED. Cette interface permet de consulter et de restaurer des documents archives, assurant la reversibilite de l'archivage conformement aux principes de conservation preconises par la norme ISO 15489:2016.")

add_figure(doc, f'{SS_DIR}/05-archives.png', 'Figure 3.19 : Interface de gestion des archives', width_cm=14)

add_heading(doc, '3.7.6 Journal d\'audit', 2)

add_para(doc, "L'interface du journal d'audit presente la liste chronologique de toutes les actions enregistrees sur la plateforme, avec des filtres par type d'action, utilisateur et date. Chaque entree affiche l'utilisateur ayant effectue l'action, le type d'action, l'entite concernee, l'adresse IP et la date. Cette interface constitue un outil essentiel pour la tracabilite, la detection d'anomalies et la conformite reglementaire, conforme aux exigences de la norme ISO 15489:2016 relative a la gestion des archives.")

add_figure(doc, f'{SS_DIR}/06-audit-log.png', 'Figure 3.20 : Journal d\'audit avec filtres', width_cm=14)

add_heading(doc, '3.7.7 Administration de la plateforme', 2)

add_para(doc, "L'interface d'administration est reservee aux super-administrateurs de la plateforme et presente une vue dediee avec une barre laterale distincte. Elle comprend les ecrans suivants : vue plateforme (statistiques globales sur les organisations, utilisateurs et documents), gestion des organisations (creation, suspension, details), gestion de la facturation (abonnements et plans), analyses (statistiques detaillees), et gestion des modules globaux.")

add_figure(doc, f'{SS_DIR}/07-administration.png', 'Figure 3.21 : Console d\'administration de la plateforme', width_cm=14)

add_figure(doc, f'{SS_DIR}/13-super-admin-organizations.png', 'Figure 3.22 : Gestion des organisations (Super Admin)', width_cm=14)

add_figure(doc, f'{SS_DIR}/15-super-admin-analytics.png', 'Figure 3.23 : Analyses de la plateforme (Super Admin)', width_cm=14)

add_heading(doc, '3.7.8 Gestion des modules', 2)

add_para(doc, "L'interface de gestion des modules presente la liste des modules disponibles pour l'organisation courante, organises par categorie. Chaque module affiche son nom, sa description, son statut (actif, inactif, suspendu) et un bouton d'activation/desactivation. Les dependances entre modules sont indiquees visuellement, et le systeme empeche l'activation d'un module si ses dependances ne sont pas satisfaites.")

add_figure(doc, f'{SS_DIR}/17-modules-page.png', 'Figure 3.24 : Gestion des modules de l\'organisation', width_cm=14)

add_heading(doc, '3.7.9 Gestion des workflows', 2)

add_para(doc, "L'interface de gestion des workflows permet de visualiser les workflows existants et d'en creer de nouveaux via un constructeur visuel. Le constructeur de workflows offre une interface intuitive pour ajouter des etats, definir des transitions entre les etats et specifier les roles autorises pour chaque transition. Les workflows existants sont affiches avec leurs etats et transitions, permettant une comprehension rapide du circuit de validation.")

add_figure(doc, f'{SS_DIR}/18-workflows-page.png', 'Figure 3.25 : Gestion des workflows', width_cm=14)

add_heading(doc, '3.7.10 Notifications et parametres', 2)

add_para(doc, "L'interface de notifications presente la liste des messages recus par l'utilisateur, avec la possibilite de les marquer comme lus. L'interface des parametres permet a l'utilisateur de configurer ses preferences, incluant les notifications par email et les options d'affichage. Ces interfaces contribuent a la personnalisation de l'experience utilisateur et au suivi des actions requises.")

add_figure(doc, f'{SS_DIR}/19-notifications-page.png', 'Figure 3.26 : Centre de notifications', width_cm=14)

add_figure(doc, f'{SS_DIR}/20-settings-page.png', 'Figure 3.27 : Parametres de l\'utilisateur', width_cm=14)

add_heading(doc, '3.7.11 Navigation adaptive', 2)

add_para(doc, "Le systeme de navigation de l'application est adapte au type d'organisation et au role de l'utilisateur. La barre laterale affiche les elements de navigation pertinents en fonction du role : un VIEWER ne voit que les Documents, les Archives et les Notifications, tandis qu'un ORG_ADMIN voit egalement l'Administration, les Modules et les Workflows. Le type d'organisation influence les labels de navigation : une universite verra 'Academique', un hopital 'Medical', un gouvernement 'Procedures'. Cette adaptation contextuelle de la navigation reduit la complexite perçue par l'utilisateur et guide naturellement vers les fonctionnalites pertinentes pour son contexte professionnel.")

add_para(doc, "Le tableau de bord de l'archiviste illustre cette adaptation : seules les sections Documents et Archives sont accessibles, avec une mise en avant des fonctionnalites d'archivage et de recherche. De meme, le tableau de bord du secretaire presente une vue simplifiee orientee vers la gestion des documents entrants et la saisie de donnees.")

add_figure(doc, f'{SS_DIR}/10-archivist-archives.png', 'Figure 3.28 : Vue archive de l\'archiviste', width_cm=14)

add_figure(doc, f'{SS_DIR}/11-secretary-dashboard.png', 'Figure 3.29 : Tableau de bord du secretaire', width_cm=14)

doc.save(OUTPUT)
print(f"Part 2 saved: {os.path.getsize(OUTPUT)/1024:.1f} KB")
